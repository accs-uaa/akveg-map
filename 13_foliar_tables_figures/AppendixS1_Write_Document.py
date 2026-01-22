# ---------------------------------------------------------------------------
# Write Appendix S1
# Author: Timm Nawrocki, Alaska Center for Conservation Science
# Last Updated: 2026-01-21
# Usage: Must be executed in a Python 3.12+ installation.
# Description: "Write Appendix S1" writes Appendix S1 from model results and plots.
# ---------------------------------------------------------------------------

# Import libraries
import os
import json
import pandas as pd
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

#### SET UP DIRECTORIES, FILES, AND FIELDS
####____________________________________________________

# Set root directory
drive = 'C:/'
root_folder = 'ACCS_Work'

# Define folder structure
document_folder = os.path.join(drive, root_folder,
                               'Projects/VegetationEcology/AKVEG_Map/Documents',
                               'Manuscript_FoliarCover_FloristicGradients')
repository_folder = os.path.join(drive, root_folder,
                                 'Repositories/akveg-map')
plot_folder = os.path.join(document_folder, 'appendix_s1/figures')
text_folder = os.path.join(repository_folder, '13_foliar_tables_figures/text')

# Define input files
template_input = os.path.join(document_folder, 'appendix_s1/Appendix_S1_template.docx')
training_input = os.path.join(document_folder, 'tables/00_Training_Data_Summary.xlsx')

# Define output files
appendix_output = os.path.join(document_folder, 'appendix_s1/Appendix_S1.docx')

# Define indicators and order
indicators = ['picgla', 'picmar', 'picsit', 'tsumer', 'bettre', 'populbt', 'poptre', 'brotre',
              'alnus', 'ndsalix', 'betshr', 'rubspe', 'vaculi', 'rhoshr', 'vacvit',
              'empnig', 'nerishr', 'dryas', 'dsalix', 'mwcalama', 'erivag', 'wetsed', 'sphagn']

#### WRITE INTRODUCTION
####____________________________________________________

# Read the training data
training_data = pd.read_excel(training_input, sheet_name='site_visits')

# Load the template
doc = Document(template_input)

# Print style names
for style in doc.styles:
    print(style.name)

# Clear existing content
for paragraph in doc.paragraphs:
    p = paragraph._element
    p.getparent().remove(p)

# Add the introductory text
doc.add_paragraph(
    'Supporting Information for “Foliar cover maps of diagnostic species predict most of the information quantity stored in vegetation types across multiple bioclimatic zones” in Ecological Monographs by Timm W. Nawrocki, Matthew J. Macander, Aaron F. Wells, Amanda Droghini, Gerald V. Frost, Lindsey A. Flagstad, Matthew L. Carlson, Calvin B. Heslop, Hunter A. Gravley, Michael Hannam, Amy E. Miller, Carl Roland, Tina Boucher, Kathryn C. Baer, Blaine T. Spellman, Marji Patz, Lisa B. Saperstein, Denise Gordon, Caitlin Willier, Elizabeth Powers.',
    style='Normal')

# Add the Heading 1
doc.add_paragraph(
    'Appendix S1: Results for Diagnostic Species Sets',
    style='Heading 1')

# Add introductory text
doc.add_paragraph(
    'For each diagnostic species set, we provide the following results:',
    style='Normal')

# Add a numbered list
outline = [
    '1. The number of site visits used to inform the map development.',
    '2. The selection rationale for the final composite model.'
    '3. A basic description of the diagnostic species set.',
    '4. A plot showing the observed versus predicted comparison from the combined test partitions of the outer cross-validation iterations for the final composite model at the site scale.',
    '5. A plot showing the observed versus predicted comparison from the combined test partitions of the outer cross-validation iterations for the final composite model at the local scale.',
    '6. The order and magnitude of the 10 most important covariates from the classifier and regressor combined into a single plot (20 covariates total).',
    '7.	A list of all constituent taxa represented by the diagnostic species set, unless this information is included in the description. Additional species may be represented but only as model error.'
]
for item in outline:
    doc.add_paragraph(item, style='List Paragraph')

# Add introductory text
doc.add_paragraph(
    'We use the term “map domain” to refer to the area mapped for this study. We use scientific names, which follow taxon concepts and synonymy in Nawrocki et al. (2024), throughout this appendix to ensure consistency across diagnostic species sets.',
    style='Normal')

#### ADD INFORMATION FOR EACH INDICATOR
####____________________________________________________

# Initialize figure count
figure_count = 1

# Add information by looping through indicators
for indicator in indicators:
    # Read json file
    json_input = os.path.join(text_folder, indicator + '.json')
    with open(json_input, 'r', encoding='utf-8') as file:
        indicator_data = json.load(file)

    # Extract values from the loaded data
    indicator_title = indicator_data.get('title')
    indictor_name = indicator_data.get('name')
    indicator_rationale = indicator_data.get('rationale')
    indicator_description = indicator_data.get('description')
    species_list = indicator_data.get('species', [])

    # Add page break
    doc.add_page_break()

    # Add the Heading 2 with parsing of italics
    h2 = doc.add_paragraph(style='Heading 2')
    heading_parts = re.split(r'(<i>.*?</i>)', indicator_title)
    for part in heading_parts:
        if part.startswith('<i>') and part.endswith('</i>'):
            # Strip the tags and format as italics
            clean_text = part.replace('<i>', '').replace('</i>', '')
            run = h2.add_run(clean_text)
            run.italic = True
        else:
            # Add as normal text
            h2.add_run(part)

    # Add the number of site visits
    site_visits = training_data.loc[training_data['indicator'] == indicator, 'site_visits'].item()
    presences = training_data.loc[training_data['indicator'] == indicator, 'presence'].item()
    p = doc.add_paragraph(style='Normal')
    run = p.add_run('Site visits:')
    run.bold = True
    run = p.add_run(f' {presences} presences out of {site_visits} site visits.')

    # Add the selection rationale
    p = doc.add_paragraph(style='Normal')
    run = p.add_run('Selection criteria:')
    run.bold = True
    run = p.add_run(f' {indicator_rationale}')

    # Add observed versus predicted plot
    performance_plot = os.path.join(plot_folder, f'figure_performance_{indicator}.png')
    doc.add_picture(performance_plot, width=Inches(6.5))

    # Add caption for observed versus predicted plot
    caption_text = f'Figure S{figure_count}. Observed foliar cover compared to predicted foliar cover for {indictor_name} at the site scale (individual sites) and local scale (sites aggregated within 10 km grids).'
    figure_count += 1
    p = doc.add_paragraph(style='Normal')
    caption_parts = re.split(r'(<i>.*?</i>)', caption_text)
    for part in caption_parts:
        if part.startswith('<i>') and part.endswith('</i>'):
            # Strip the tags and format as italics
            clean_text = part.replace('<i>', '').replace('</i>', '')
            run = p.add_run(clean_text)
            run.italic = True
        else:
            # Add as normal text
            p.add_run(part)

    # Add the description with parsing of italics
    p = doc.add_paragraph(style='Normal')
    description_parts = re.split(r'(<i>.*?</i>)', indicator_description)
    for part in description_parts:
        if part.startswith('<i>') and part.endswith('</i>'):
            # Strip the tags and format as italics
            clean_text = part.replace('<i>', '').replace('</i>', '')
            run = p.add_run(clean_text)
            run.italic = True
        else:
            # Add as normal text
            p.add_run(part)

    # Add covariate importance plot
    covariate_plot = os.path.join(plot_folder, f'figure_importance_{indicator}.png')
    doc.add_picture(covariate_plot, width=Inches(6.5))

    # Add caption for covariate plot
    caption_text = f'Figure S{figure_count}. Ten most important covariates from the classifier and regressor for {indictor_name} (for a total of 20 covariates). Covariates are named by the abbreviations provided in Appendix S4.'
    figure_count += 1
    p = doc.add_paragraph(style='Normal')
    caption_parts = re.split(r'(<i>.*?</i>)', caption_text)
    for part in caption_parts:
        if part.startswith('<i>') and part.endswith('</i>'):
            # Strip the tags and format as italics
            clean_text = part.replace('<i>', '').replace('</i>', '')
            run = p.add_run(clean_text)
            run.italic = True
        else:
            # Add as normal text
            p.add_run(part)

    # Add constituent taxa section if it exists
    if species_list != []:

        # Add constituent taxa header
        doc.add_paragraph('Constituent Taxa', style='Heading 3')

        # Add a continuous section break
        current_section = doc.add_section(WD_SECTION.CONTINUOUS)

        # Configure two columns
        sectPr = current_section._sectPr
        cols = sectPr.xpath('./w:cols')

        # Get or create the columns element
        if not cols:
            cols = OxmlElement('w:cols')
            sectPr.append(cols)
        else:
            cols = cols[0]

        # Set the number of columns to '2'
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720')

        # Add species list
        for taxon in species_list:
            p = doc.add_paragraph(style='Normal')
            # Align left (overrides justified in template)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Apply hanging indent
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            # Write line
            taxon_parts = re.split(r'(<i>.*?</i>)', taxon)
            for part in taxon_parts:
                if part.startswith('<i>') and part.endswith('</i>'):
                    # Strip the tags and format as italics
                    clean_text = part.replace('<i>', '').replace('</i>', '')
                    run = p.add_run(clean_text)
                    run.italic = True
                else:
                    # Add as normal text
                    p.add_run(part)

        # Add a continuous section break to end the columns
        new_section_single = doc.add_section(WD_SECTION.CONTINUOUS)
        sectPr_single = new_section_single._sectPr
        cols_single = sectPr_single.xpath('./w:cols')

        # Safely get or create the columns element for the new section
        if not cols_single:
            cols_single = OxmlElement('w:cols')
            sectPr_single.append(cols_single)
        else:
            cols_single = cols_single[0]

        # Reset columns to '1'
        cols_single.set(qn('w:num'), '1')

#### ADD LITERATURE CITED
####____________________________________________________

# Add page break
doc.add_page_break()

# Read references json file
json_input = os.path.join(text_folder, 'references.json')
with open(json_input, 'r', encoding='utf-8') as file:
    references_data = json.load(file)

# Extract values from the loaded data
references = references_data.get('references')

# Add references
doc.add_paragraph('References', style='Heading 2')

# Write references
for reference in references:
    p = doc.add_paragraph(reference, style='Normal')
    # Align left
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Apply hanging indent
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)

#### EXPORT DOCUMENT
####____________________________________________________

# Save the new document
doc.save(appendix_output)
